#last updates : 9/08/2022,11am
from pdf2docx import parse #pip install pdf2docx
import fitz # pip install PyMuPDF
from docx import Document  #pip install docx
from os import stat,remove,path
"""
    this module has 4 function that can be used to convert:
    a pdf file to document
        pdf_to_doc(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
    a pdf file to text
        pdf_to_text(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
    a text file to document
        text_to_doc(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
    a doc file to text    
        doc_to_text(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
"""

#custom error
class CustomError(Exception):
    error_data=""
    
#function that returns file format
def file_type(input_file:str):    
    file_extension=path.splitext(input_file)
    return(file_extension[1].lower())

#function that returns the size of the file in MB
def filesize(input_file:str):
    return (path.getsize(input_file)/1048576)
#------function 1----------
def pdf_to_doc(input_file:str,output_file:str):
    """
        takes a pdf file as input and converts it into a doc file 
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:
        #checking for correct file format
        input_file_types=['.pdf']
        output_file_types=[".docx",".doc"]
        input_file_type=file_type(input_file)
        if input_file_type not in input_file_types:
            CustomError.error_data="Given file format ("+input_file_type+") is not supported..!"
            raise CustomError
        output_file_type=file_type(output_file)
        if output_file_type not in output_file_types:
            CustomError.error_data="Given file format ("+output_file_type+") is not supported..!"
            raise CustomError

        #cheking if the file is too big
        max_size=10#confirm the max size
        if(filesize(input_file)>max_size):
            CustomError.error_data="Maximum file size is "+str(max_size)+"MB."
            raise CustomError

        pages=()
        #Converts pdf to docx
        if pages:
            #iterates through pdf file and extracts each page and write it on doc file
            pages = [int(i) for i in list(pages) if i.isnumeric()]
        result = parse(input_file,output_file,pages=pages)
        
        #cheking if the file is empty
        try:
        #document object for docx file
            document= Document(output_file)
        except:
            document=Document()
            document.save(output_file)
        lines=document.paragraphs
        isempty=""
        for data in lines:
            isempty+=data.text
            if(isempty!=""):
                break
        if(isempty=="" or isempty==" "):
            CustomError.error_data="file is empty or may be password protected"
            raise CustomError
        
        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except CustomError:
        #removing  previously created output file
        #remove(output_file)
        #writing error
        return_data["error"]=CustomError.error_data
    except Exception as e:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data['error']=e  
    return return_data   

#------function 2----------
def pdf_to_text(input_file:str,output_file:str):
    """
        takes a pdf file as input and converts it into a text file
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:
        #checking for correct file format
        input_file_types=[".pdf"]
        output_file_types=[".txt"]
        input_file_type=file_type(input_file)
        if input_file_type not in input_file_types:
            CustomError.error_data="Given file format ("+input_file_type+") is not supported..!"
            raise CustomError
        output_file_type=file_type(output_file)
        if output_file_type not in output_file_types:
            CustomError.error_data="Given file format ("+output_file_type+") is not supported..!"
            raise CustomError

        #cheking if the file is too big
        max_size=10#confirm the max size
        if(filesize(input_file)>max_size):
            CustomError.error_data="Maximum file size is "+str(max_size)+"MB."
            raise CustomError

        #open the pdf file as doc variable
        with fitz.open(input_file) as doc:
            text = ""
            #open output file for writitng
            with open(output_file,'w', encoding="utf-8") as text_file:
                #each page in pdf file
                for page in doc:
                    #data from each page is retrived using get_text() and written on output file
                    text_file.write(page.get_text())
        print(stat(output_file).st_size,type(stat(output_file).st_size))
        #cheking if the file is empty
        if(stat(output_file).st_size <= 3):
            CustomError.error_data="file is empty..!"
            raise CustomError

        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except CustomError:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data["error"]=CustomError.error_data
    except ValueError:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data["error"]="file is password protected"
    except Exception as e:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data['error']=e  
    return return_data   


#------function 3----------
def text_to_doc(input_file:str,output_file:str):
    """
        takes a text file as input and converts it into a doc file 
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try: 
        #checking for correct file format
        input_file_types=[".txt"]
        output_file_types=[".docx",".doc"]
        input_file_type=file_type(input_file)
        if input_file_type not in input_file_types:
            CustomError.error_data="Given file format ("+input_file_type+") is not supported..!"
            raise CustomError
        output_file_type=file_type(output_file)
        if output_file_type not in output_file_types:
            CustomError.error_data="Given file format ("+output_file_type+") is not supported..!"
            raise CustomError

        #cheking if the file is too big
        max_size=10#confirm the max size
        if(filesize(input_file)>max_size):
            CustomError.error_data="Maximum file size is "+str(max_size)+"MB."
            raise CustomError

        #cheking if the file is empty
        if stat(input_file).st_size == 0:
            CustomError.error_data="file is empty"
            raise CustomError
        
        with open(input_file) as text_file:
            #reading lines in the txt file as list
            lines=text_file.readlines()
            #document object
            document_file=Document()
            for data in lines:
                #writing text on docx file
                document_file.add_paragraph(data)
                document_file.save(output_file) 
        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except CustomError:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data["error"]=CustomError.error_data
    except Exception as e:
        remove(output_file)
        #writing error
        return_data['error']=e  
    return return_data   


#------function 4----------
def doc_to_text(input_file:str,output_file:str):
    """
        takes a doc file as input and converts it into a text file 
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:
        #checking for correct file format
        input_file_types=[".docx",".doc"]
        output_file_types=[".txt"]
        input_file_type=file_type(input_file)
        if input_file_type not in input_file_types:
            CustomError.error_data="Given file format ("+input_file_type+") is not supported..!"
            raise CustomError
        output_file_type=file_type(output_file)
        if output_file_type not in output_file_types:
            CustomError.error_data="Given file format ("+output_file_type+" is not supported..!"
            raise CustomError
        
        #cheking if the file is too big
        max_size=10#confirm the max size
        if(filesize(input_file)>max_size):
            CustomError.error_data="Maximum file size is "+str(max_size)+"MB."
            raise CustomError
        try:
        #document object for docx file
            document= Document(input_file)
        except:
            document=Document()
            document.save(input_file)
        #finding total lines in document
        #document.paragraphs returns a list type which contains each line in doc file as a list element
        lines=document.paragraphs

        #writing on txt file
        with open(output_file,"w") as text_file:
            for data in lines:
                text_file.write(data.text+"\n")
            
        #cheking if the file is empty
        if(stat(output_file).st_size <=3):
            CustomError.error_data="file is empty"
            raise CustomError

        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except CustomError:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data["error"]=CustomError.error_data
    except Exception as e:
        #removing  previously created output file
        remove(output_file)
        #writing error
        return_data['error']=e  
    return return_data   