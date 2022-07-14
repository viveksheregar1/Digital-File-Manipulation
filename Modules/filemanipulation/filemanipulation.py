#last updates : 14/07/2022
from pdf2docx import parse #pip install pdf2docx
import fitz # pip install PyMuPDF
import docx #pip install docx
"""
    this module has 4 function that can be used to convert:
    
    a pdf file to document
        pdf_to_doc(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
    
    a pdf file to text
        pdf_to_text(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
    
    a text file to document
        text_to_doc(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
         
    a doc file to text    
        doc_to_text(input_file:str,output_file:str)
            returns a dictionary return_data{'output':"",'error':"","output_file":""}
"""
#------function 1----------
#in production
def pdf_to_doc(input_file:str,output_file:str):
    #in production for (text from image for chandrakala mam)
    """
        takes a pdf file as input and converts it into a doc file 
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:
        pages=()
        #Converts pdf to docx
        if pages:
            #iterates through pdf file and extracts each page and write it on doc file
            pages = [int(i) for i in list(pages) if i.isnumeric()]
        result = parse(input_file,output_file, pages=pages)
        print("done")
        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except Exception as e:
        #writing error
        return_data['error']=e  
    return return_data   

#------function 2----------
def pdf_to_text(input_file:str,output_file:str):
    """
        takes a pdf file as input and converts it into a text file
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:
        #open the pdf file as doc variable
        with fitz.open(input_file) as doc:
            text = ""
            #open output file for writitng
            with open(output_file,'w', encoding="utf-8") as text_file:
                #each page in pdf file
                for page in doc:
                    #data from each page is retrived using get_text() and written on output file
                    text_file.write(page.get_text())
        #print("done")
        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except Exception as e:
        #writing error
        return_data['error']=e  
    return return_data   


#------function 3----------
def text_to_doc(input_file:str,output_file:str):
    """
        takes a text file as input and converts it into a doc file 
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:  
        with open(input_file) as text_file:
            #reading lines in the docx file as list
            lines=text_file.readlines()
            #document object
            document_file=docx.Document()
            for data in lines:
                #writing text on docx file
                document_file.add_paragraph(data)
                document_file.save(output_file) 
        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except Exception as e:
        #writing error
        return_data['error']=e  
    return return_data   


#------function 4----------
def doc_to_text(input_file:str,output_file:str):
    """
        takes a doc file as input and converts it into a text file 
    """
    #dictionary to store output 
    return_data={
        "output":"",
        "error":"",
        "output_file":""
    }
    try:
        #first way
        #passing docx file
        #data=docx2txt.process(input_file)
        #saving content extracted from docx as data to outputdata.txt
        #with open(output_file,"w") as text_file:
        #   print(data,file=text_file)
        #   print('complete')"""

        #second way
        #document object for docx file
        document= docx.Document(input_file)

        #finding total lines in document
        #document.paragraphs returns a list type which contains each line in doc file as a list element
        lines=document.paragraphs

        #writing on txt file
        with open(output_file,"w") as text_file:
            for data in lines:
                text_file.write(data.text+"\n")
        #storing output in dictionary
        return_data['output']="done"
        return_data['output_file']=output_file #path of saved file
    except Exception as e:
        #writing error
        return_data['error']=e  
    return return_data   

input_file=r"C:\Users\vivek\OneDrive\Desktop\College Project\modules\filemanipulation\txt_file.txt"
output_file=r"C:\Users\vivek\OneDrive\Desktop\College Project\modules\filemanipulation\output.pdf"
print((input_file,output_file))
