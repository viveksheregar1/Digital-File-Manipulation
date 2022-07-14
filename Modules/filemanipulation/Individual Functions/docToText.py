import docx
#import docx2txt
#first way
#passing docx file
#data=docx2txt.process(r"C:\Users\vivek\OneDrive\Desktop\College Project\modules\file manipulation\testdoc.docx")
#saving content extracted from docx as data to outputdata.txt
#with open(r"C:\Users\vivek\OneDrive\Desktop\College Project\modules\file manipulation\outputdata.txt","w") as text_file:
#   print(data,file=text_file)
#   print('complete')"""

#second way
#document object for docx file
document= docx.Document(r"C:/Users/vivek/OneDrive/Desktop/College Project/modules/fileManipulation/doc_file.docx")

#finding total lines in document
#document.paragraphs returns a list type which contains each line in doc file as a list element
lines=document.paragraphs

#writing on txt file
with open(r"C:/Users/vivek/OneDrive/Desktop/College Project/modules/fileManipulation/doc_to_text.txt","w") as text_file:
    for data in lines:
        text_file.write(data.text+"\n")
    print('complete')
